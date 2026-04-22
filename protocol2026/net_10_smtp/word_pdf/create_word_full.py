import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 加粗
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸
from docx.shared import RGBColor  # 颜色模块
import sys
from pathlib import Path

# 获取当前文件所在目录的父目录（项目根目录）并添加到Python路径
current_file = Path(__file__)
current_dir = current_file.parent
project_root = current_file.parent.parent.parent
sys.path.append(str(project_root))

from net_10_smtp.modules.mat_bing import mat_bing
import os

"""
----------------Windows----------------
1. python-docx
2. MS Office
3. 字体默认有
4. docx2pdf


----------------Linux----------------
1. python-docx
2. yum install libreoffice
3. cd /usr/share/fonts
   sudo wget https://noto-website-2.storage.googleapis.com/pkgs/NotoSansCJKsc-hinted.zip
   sudo unzip NotoSansCJKsc-hinted.zip -d noto
   sudo fc-cache -fv
4. os.popen(f"libreoffice --headless "
            f"--convert-to pdf "
            f"--outdir ./word_pdf/saved_pdf ./word_pdf/saved_word/syslog.docx")
"""


def create_word_full(student_data, img_counters, img_protocols, save_word_name):
    today = time.strftime("%Y{y}%m{m}%d{d}", time.localtime()).format(y='年', m='月', d='日')

    document = Document()

    document.styles['Normal'].font.name = 'Noto Sans CJK SC'
    document.styles['Normal'].font.size = Pt(14)
    # 设置文档的基础字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    # 设置文档的基础样式

    document.add_picture(f'{current_dir}{os.sep}src_img{os.sep}logo.png', width=Inches(6))
    # 在文件最上头插入图，宽度为6英寸

    # 初始化建立第一个自然段
    p1 = document.add_paragraph()
    # 对齐方式为居中，没有这句的话默认左对齐。
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 写入内容'乾颐堂网络技术服务协议'
    run1 = p1.add_run('乾颐堂网络技术服务协议')

    # 设置字体font
    run1.font.name = 'Noto Sans CJK SC'
    # 设置中文字体
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')

    # 设置字体大小为20磅
    run1.font.size = Pt(20)
    # 设置加粗
    run1.font.bold = True
    # 段后距离5磅
    p1.space_after = Pt(1)
    # 段前距离5磅
    p1.space_before = Pt(5)

    # 起一个新的段落
    p2 = document.add_paragraph()

    # 写入内容 "报名日期：2022年12月03日"
    run2 = p2.add_run('报名日期：%s' % (today,))

    # 调整字体与大小
    run2.font.name = 'Noto Sans CJK SC'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    run2.font.size = Pt(12)
    run2.font.bold = True

    # 创建一个 7行 6列 的表
    table = document.add_table(rows=7, cols=6, style='Light Grid Accent 2')
    # table_run1 = table.cell(0, 0).paragraphs[0].add_run('姓名')
    # table.cell(0,1).text= Trainee['姓名']
    # table.cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # table.font.name = u'隶书'
    # table._element.rPr.rFonts.set(qn('w:eastAsia'), u'隶书')

    # 写入表格内容
    table.cell(0, 0).text = '姓名'
    table.cell(0, 1).text = student_data['姓名']
    table.cell(0, 2).text = '性别'
    table.cell(0, 3).text = student_data['性别']
    table.cell(0, 4).text = '出生年月'
    table.cell(0, 5).text = student_data['出生年月']
    table.cell(1, 0).text = '籍贯'
    table.cell(1, 1).text = student_data['籍贯']
    table.cell(1, 2).text = '学历'
    table.cell(1, 3).text = student_data['学历']
    table.cell(1, 4).text = '单位/学校'
    table.cell(1, 5).text = student_data['单位/学校']
    table.cell(2, 0).text = 'QQ'
    table.cell(2, 1).text = student_data['QQ']
    table.cell(2, 4).text = '联系电话'
    table.cell(2, 5).text = student_data['联系电话']
    table.cell(3, 0).text = '身份证号'
    # 身份证号占用 1 - 3 三列
    table.cell(3, 1).text = student_data['身份证号']
    table.cell(3, 4).text = '从何得知乾颐堂'
    table.cell(3, 5).text = student_data['从何得知乾颐堂']
    table.cell(4, 0).text = '报名课程'
    table.cell(4, 1).text = student_data['报名课程']
    table.cell(4, 4).text = '税费金额（不含发票税）'
    table.cell(4, 5).text = student_data['税费金额（不含发票税）']
    table.cell(5, 0).text = '课程顾问'
    table.cell(5, 1).text = student_data['课程顾问']
    table.cell(5, 4).text = '班主任联系电话'
    table.cell(5, 5).text = student_data['班主任联系电话']
    table.cell(6, 0).text = '备注'
    table.cell(6, 1).merge(table.cell(6, 5))  # 合并单元格
    table.cell(4, 1).merge(table.cell(4, 3))  # 合并单元格
    table.cell(2, 1).merge(table.cell(2, 2))  # 合并单元格
    table.cell(3, 1).merge(table.cell(3, 3))  # 合并单元格
    table.cell(5, 1).merge(table.cell(5, 3))  # 合并单元格
    table.cell(2, 1).merge(table.cell(2, 3))  # 合并单元格

    # 开始全新的一页
    document.add_page_break()  # 添加分页符

    # 起一个新的段落
    p3 = document.add_paragraph()

    # 对齐方式为居中，没有这句的话默认左对齐。
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 写入内容 '关于乾颐堂科技有限责任公司介绍'
    run3 = p3.add_run('关于乾颐堂科技有限责任公司介绍')

    # 调整字体与颜色
    run3.font.name = 'Noto Sans CJK SC'
    run3.font.color.rgb = RGBColor(0xaf, 0x26, 0x26)

    # 起一个新的段落
    p5 = document.add_paragraph()
    # 写入内容
    run5 = p5.add_run('    北京乾颐堂科技有限责任公司于2014年9月24日成立，上海乾颐堂成立时间是2014年9月7日，'
                      '北京负责人是现任明教教主-秦柯，上海负责人是Ender安德-周亚军。乾颐堂在北京、上海、'
                      '南京均设立了专业的网络实验室，我们创建了一个专注于网络技术、Cisco认证考试以及华为HCIE认证考试的培训机构。'
                      '在很短的时间内，学员已经达到5000多人，我们在口碑以及培训水平上远远超过其他同类的培训中心。')
    # 调整字体与文字大小
    run5.font.name = 'Noto Sans CJK SC'
    run5._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    run5.font.size = Pt(10)

    # 起一个新的段落
    p5 = document.add_paragraph()
    # 写入内容
    run5 = p5.add_run('    乾颐堂师资力量雄厚，包括现任明教教主，安德老师等数十位重量级的老师。我们的理念是为您想的更多，'
                      '旨在将高中低端各个层级的IT认证培训做得更加专业更加多样化，力争成为行业内教学质量第一、服务质量第一的培训机构。')
    # 调整字体与文字大小
    run5.font.name = 'Noto Sans CJK SC'
    run5._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    run5.font.size = Pt(10)

    # ------------------------测试字体与大小部分---------------------
    # 设置粗体字
    document.add_paragraph('设置粗体字:').add_run('粗体字').bold = True
    # 设置斜体字
    document.add_paragraph('设置斜体字:').add_run('斜体字').italic = True
    # 设置字号50
    document.add_paragraph('设置字号50:').add_run('50').font.size = Pt(50)
    # 设置字体颜色为 af2626
    document.add_paragraph('设置字体颜色:').add_run('颜色').font.color.rgb = RGBColor(0xaf, 0x26, 0x26)
    # 样式叠加: 将字体改到30号并且将字体改成特定颜色;
    run5 = document.add_paragraph('同时设置文字颜色和字号:').add_run('颜色和尺寸')
    run5.font.size = Pt(30)
    run5.font.color.rgb = RGBColor(0xa, 0x26, 0x26)

    # 创建 有序列表
    document.add_paragraph('').add_run('有序列表').font.size = Pt(30)
    document.add_paragraph('cisco security', style='List Number')
    document.add_paragraph('cisco wireless', style='List Number')
    document.add_paragraph('python', style='List Number')

    # 创建 无序列表
    document.add_paragraph('').add_run('无序列表').font.size = Pt(30)
    document.add_paragraph('天地匆匆 惊鸿而过 路有千百个', style='List Bullet')
    document.add_paragraph('遑遑无归 闲云逸鹤 人间红尘过', style='List Bullet')
    document.add_paragraph('引势而流 鸿门乱局 各有各选择', style='List Bullet')
    document.add_paragraph('乾震坎艮 坤巽离兑 定一切生克', style='List Bullet')

    # 开始全新的一页
    document.add_page_break()  # 添加分页符
    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('乾颐堂学员报名课程分布如下图')
    run4.font.name = 'Noto Sans CJK SC'

    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    run4.font.size = Pt(16)
    run4.font.bold = True

    # 产生图片
    mat_bing(img_counters, img_protocols, 'temp.png')
    # 插入图片
    document.add_picture('temp.png', width=Inches(5.0), height=Inches(5.0))
    # 删除零时产生的图片
    os.remove('temp.png')

    # 保持产生的word文件
    document.save(save_word_name)


if __name__ == '__main__':
    data = {'姓名': '张三', '性别': '男', '出生年月': '1997.6.2', '籍贯': '北京', '学历': '本科',
            '单位/学校': 'IC PARK',
            'QQ': '8888888', '联系电话': '13999999999', '身份证号': '156456369263561666',
            '从何得知乾颐堂': '老学员介绍',
            '报名课程': 'python基础', '税费金额（不含发票税）': '499', '课程顾问': '小雪',
            '班主任联系电话': '13432555669', '备注:': ''}
    counters = [30, 53, 12, 45, 100]
    protocols = ['vip学员', '安全学员', '无线学员', 'python学员', 'NetDevOps']

    create_word_full(data, counters, protocols, f'{current_dir}{os.sep}saved_word{os.sep}python-docx演示.docx')
