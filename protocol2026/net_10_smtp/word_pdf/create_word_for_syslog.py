from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 加粗
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸
from docx.shared import RGBColor  # 颜色模块
import sys
from pathlib import Path

# 获取当前文件所在目录的父目录（项目根目录）并添加到Python路径
current_file = Path(__file__)
current_dir = current_file.parent
project_root = current_file.parent.parent.parent
sys.path.append(str(project_root))

from net_10_smtp.modules.syslog_bing import syslog_bing
import os


def create_word_for_syslog(add_img, save_word_name):
    document = Document()

    # 设置文档的基础字体
    document.styles['Normal'].font.name = 'Noto Sans CJK SC'
    document.styles['Normal'].font.size = Pt(14)

    # 设置文档的基础样式
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    # 在文件最上头插入图，宽度为6英寸
    document.add_picture(add_img, width=Inches(6))

    # 初始化建立第一个自然段
    p1 = document.add_paragraph()
    # 对齐方式为居中，没有这句的话默认左对齐。
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置标题
    run1 = p1.add_run('乾颐堂Python强化班Syslog分析')
    # 设置字体font
    run1.font.name = 'Noto Sans CJK SC'
    # 设置中文字体
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    # 设置字体大小为20磅
    run1.font.size = Pt(20)
    # 设置加粗
    run1.font.bold = True
    # 段后距离1磅
    p1.space_after = Pt(1)
    # 段前距离5磅
    p1.space_before = Pt(5)

    # 初始化建立第二个自然段
    p2 = document.add_paragraph()

    # 产生数据和图
    syslog_result = syslog_bing('temp.png')
    # 第二个自然段主题
    run2 = p2.add_run('下面是最近一个小时的Syslog的数据统计! 显示排前三的Syslog严重级别与数量')
    # 字体和大小
    run2.font.name = 'Noto Sans CJK SC'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    run2.font.size = Pt(10)

    # 创建表
    table = document.add_table(rows=4, cols=3, style='Light Grid Accent 2')

    # 第一行
    table.cell(0, 0).text = '严重级别'
    table.cell(0, 1).text = '数量'
    table.cell(0, 2).text = '百分比'

    # print(syslog_result)
    total = sum([y for x, y in syslog_result])

    # 后续行
    i = 1
    for x, y in syslog_result[:3]:
        # print(x, y)
        table.cell(i, 0).text = x
        table.cell(i, 1).text = str(int(y))
        table.cell(i, 2).text = f'{(y/total)*100:.1f}'
        i += 1

    # 初始化建立第三个自然段
    p3 = document.add_paragraph()
    # 第二个自然段主题
    run3 = p3.add_run('\r\n下面是最近一个小时的Syslog的数据统计饼状图分析!')
    # 字体和大小
    run3.font.name = 'Noto Sans CJK SC'
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    run3.font.size = Pt(10)
    # 插入图片
    document.add_picture('temp.png', width=Inches(3.0), height=Inches(3.0))
    # 删除图片
    os.remove('temp.png')

    # 保存文档
    document.save(save_word_name)


if __name__ == '__main__':
    create_word_for_syslog(f'{current_dir}{os.sep}src_img{os.sep}logo.png', 
                           f'{current_dir}{os.sep}saved_word{os.sep}syslog.docx')
